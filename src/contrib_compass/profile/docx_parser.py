"""
contrib_compass.profile.docx_parser — Extract raw text from DOCX files.

Responsibility:
    Given DOCX bytes, return a single string of all readable paragraph text.

NOT responsible for:
    - Parsing or interpreting the text (that's skill_normalizer's job)
    - Handling PDF or other formats (see pdf_parser)

Dependencies:
    python-docx — https://python-docx.readthedocs.io/
"""

from __future__ import annotations

import io
import logging

from docx import Document  # type: ignore[import-untyped]
from docx.oxml.ns import qn  # type: ignore[import-untyped]

logger = logging.getLogger(__name__)


class DOCXParseError(ValueError):
    """Raised when the DOCX cannot be opened or has no extractable text."""


def extract_text(docx_bytes: bytes) -> str:
    """Extract all text from a DOCX file given its raw bytes.

    Reads every paragraph (including those inside tables) and joins them with
    newlines.  Empty paragraphs are skipped.

    Args:
        docx_bytes: Raw binary content of a .docx file.

    Returns:
        A single string containing all readable text from the document.

    Raises:
        DOCXParseError: If the bytes are not a valid DOCX or contain no text.

    Example:
        >>> with open("resume.docx", "rb") as f:
        ...     text = extract_text(f.read())
        >>> "Python" in text
        True
    """
    if not docx_bytes:
        raise DOCXParseError("DOCX bytes are empty")

    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        raise DOCXParseError(f"Cannot open DOCX: {exc}") from exc

    lines: list[str] = []

    # ── Body paragraphs ────────────────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # ── Text inside tables (common in resumes) ─────────────────────────────
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)

    if not lines:
        raise DOCXParseError("DOCX contains no extractable text")

    return "\n".join(lines)
